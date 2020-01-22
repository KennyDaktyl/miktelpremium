#!/usr/bin/env python
# -*- coding: utf-8 -*-

import json
import coreapi
import time
import os

import win32print
# from win32print import GetDefaultPrinter,GetPrinter

from docx import Document
from docx.shared import Cm, Pt, RGBColor

pizzeria = 1

RODZAJ_DOSTAWY = (
    (1, "Lokal"),
    (2, "Wynos"),
    (3, "Dostawa"),
)
PŁATNOSC = (
    (1, "Gotówka"),
    (2, "Karta"),
    (3, "Online_1"),
    (4, "Online_2"),
    (5, "Online_3"),
)

client = coreapi.Client()

p = True
while p:
    try:
        # schema = client.get("http://localhost:8000/ws/order")
        schema = client.get("http://pizzeriasystem.herokuapp.com/ws/order")
        auth = coreapi.auth.TokenAuthentication(
            scheme='JWT',
            token='4wc6wsjzrep!+!9ttm+einr$ndb=9ic7-*ftdi!tlc5-$598sc')
        # orders = client.get("http://localhost:8000/ws/order")
        orders = client.get("http://pizzeriasystem.herokuapp.com/ws/order")
        print("Połączenie z serwerem OK")
        # print(orders)
    except:
        orders = []
        print("Brak połączenia z serwerem")
    document = Document()
    p = document.add_paragraph()

    # document.header=None
    section = document.sections[0]
    section.page_width = Cm(8)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    # from docx.enum.text import WD_LINE_SPACING

    # paragraph = document.add_paragraph()
    # paragraph_format = paragraph.paragraph_format
    # paragraph_format.left_indent = Cm(0)
    # run = document.add_paragraph().add_run()
    # font = run.font
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

    if len(orders) > 0:
        for el in orders:
            if str(el['workplace_id']['id']) == str(pizzeria):
                id_order = el['id']
                workplace_str = str(el['workplace_id']['workplace_name'])
                runner = p.add_run(str(workplace_str)+"\n")
                runner.bold = True
                number = ("Nr: " + str(el['number'])+"\n")
                p.add_run(str(number))
                date = ("Data: " + str(el['date'])+"\n")
                p.add_run(str(date))
                time_start = (
                    "Godzina: " + str(el['time_start'][0:5])+"\n")
                p.add_run(str(time_start))
                barman = (
                    "Przyjmujący: " + str(el['barman_id']['username'])+"\n")
                p.add_run(str(barman))
                type_of_order = el['type_of_order']
                pay_method = el['pay_method']
                type_order = (
                    RODZAJ_DOSTAWY[type_of_order-1][1])
                p.add_run(""+"\n")
                runner = p.add_run("Rodzaj dost.: "+str(type_order)+"\n")
                runner.bold = True
                p.add_run("---------------------- "+"\n")
                if el['address'] != None:
                    address = (
                        "Adres: " + str(el['address']['street']))
                    runner = p.add_run(address+"\n")
                    runner.bold = True
                    p.add_run(""+"\n")
                    phone = (
                        "Telefon: " + str(el['address']['client_id']['phone_number']))
                    runner = p.add_run(phone+"\n")
                    runner.bold = True
                    p.add_run(""+"\n")
                    if str(el['address']['client_id']['info']) != "":
                        text = str(el['address']['client_id']['info'])
                        p.add_run("Info: "+text+"\n")
                if el['discount'] != 0:
                    discount = ("Rabat.: " + str(el['discount'])+"%")
                    p.add_run(discount+"\n")
                    p.add_run(""+"\n")
                price = ("Kwota: " + str(el['order_total_price2']))
                runner = p.add_run(price+"\n")
                runner.bold = True
                p.add_run(""+"\n")
                runner = p.add_run(PŁATNOSC[pay_method-1][1]+"\n")
                runner.bold = True
                p.add_run(""+"\n")
                time_zero = (
                    "Godzina dos.: " + str(el['time_zero'][0:5]))
                runner = p.add_run(time_zero+"\n")
                runner.bold = True
                p.add_run(""+"\n")
                if el['info'] != None:
                    info = ("Info.: " + str(el['info']))
                    p.add_run(info+"\n")

                try:
                    # url = "http://localhost:8000/pos/"+str(id_order)
                    url = "http://pizzeriasystem.herokuapp.com/pos/" + \
                        str(id_order)
                    pos = client.get(url)
                    print("Połączenie z serwerem OK")
                except:
                    pos = []
                    print("Brak połączenia z serwerem")
                if len(pos) > 0:
                    document.add_paragraph("---------------- ")
                    document.add_paragraph("Produkty: ")
                    for el in pos:
                        if el['product_id']['category']['category_number'] == 1 or el['product_id']['category']['category_number'] == 2:
                            # print(el['product_id']['category'])
                            if el['pizza_half'] != True:
                                product = (str(el['quantity'])+"x "+str(el['product_id']['product_name'])+", "+str(
                                    el['size_id']['size_name'])+"\n"+"cena: "+str(el['price'])+", dodatki: "+str(el['extra_price']))
                                runner = document.add_paragraph().add_run(product+"\n")
                                runner.bold = True
                                if el['cake_info'] != "":
                                    cake = el['cake_info']
                                    cake = cake.replace("\r\n", "")
                                    cake = cake.replace("  ", "")
                                    document.add_paragraph(
                                        "Ciasto: "+cake+"zł"+"\n(cena wliczona w dodatki)")
                                if el['discount'] != 0:
                                    document.add_paragraph(
                                        "Rabat.: " + str(el['discount'])+"%")
                                if el['change_topps_info'] != "":
                                    changes = el['change_topps_info']
                                    changes = changes.replace("\r\n", "")
                                    changes = changes.replace("  ", "")
                                    runner = document.add_paragraph().add_run("Zmiany: ")
                                    runner.bold = True
                                    document.add_paragraph(str(changes)+"\n")
                                if el['sauces_free'] != "":
                                    sauces_free = el['sauces_free']
                                    document.add_paragraph(
                                        "Sosy: "+str(sauces_free))
                                if el['sauces_pay'] != "":
                                    sauces_pay = el['sauces_pay']
                                    sauces_pay = sauces_pay.replace("\r\n", "")
                                    sauces_pay = sauces_pay.replace("  ", "")
                                    document.add_paragraph(
                                        "Sosy płatne: "+str(sauces_pay))
                                if el['info'] != "":
                                    document.add_paragraph(
                                        "Info.: " + str(el['info']))
                                document.add_paragraph(
                                    "Cena: "+str(el['total_price']))
                                document.add_paragraph("---------------- ")
                            else:
                                product = (str(el['quantity'])+"x "+str(el['halfpizza_name'])+", "+str(
                                    el['size_id']['size_name'])+"\n"+"cena: "+str(el['price'])+", dodatki: "+str(el['extra_price']))
                                runner = document.add_paragraph().add_run(product+"\n")
                                runner.bold = True
                                if el['cake_info'] != "":
                                    cake = el['cake_info']
                                    cake = cake.replace("\r\n", "")
                                    cake = cake.replace("  ", "")
                                    document.add_paragraph(
                                        "Ciasto: "+cake+"zł"+"\n")
                                if el['discount'] != 0:
                                    document.add_paragraph(
                                        "Rabat.: " + str(el['discount'])+"%")
                                if el['change_topps_info'] != "":
                                    changes = el['change_topps_info']
                                    changes = changes.replace("\r\n", "")
                                    changes = changes.replace("  ", "")
                                    runner = document.add_paragraph().add_run("Zmiany lewa: ")
                                    runner.bold = True
                                    document.add_paragraph(str(changes)+"\n")

                                if el['change_topps_info_other_side'] != "":
                                    changes = el['change_topps_info_other_side']
                                    changes = changes.replace("\r\n", "")
                                    changes = changes.replace("  ", "")
                                    runner = document.add_paragraph().add_run("Zmiany prawa: ")
                                    runner.bold = True
                                    document.add_paragraph(str(changes)+"\n")
                                if el['sauces_free'] != "":
                                    sauces_free = el['sauces_free']
                                    document.add_paragraph(
                                        "Sosy: "+str(sauces_free))
                                if el['sauces_pay'] != "":
                                    sauces_pay = el['sauces_pay']
                                    sauces_pay = sauces_pay.replace("\r\n", "")
                                    sauces_pay = sauces_pay.replace("  ", "")
                                    document.add_paragraph(
                                        "Sosy płatne: "+str(sauces_pay))
                                if el['info'] != "":
                                    document.add_paragraph(
                                        "Info.: " + str(el['info']))
                                document.add_paragraph(
                                    "Cena: "+str(el['total_price']))
                                document.add_paragraph("---------------- ")
                        else:
                            product = (str(
                                el['quantity'])+"x "+str(el['product_id']['product_name'])+"\n"+"cena: "+str(el['price']))
                            runner = document.add_paragraph().add_run(product+"\n")
                            runner.bold = True
                            if el['extra_price'] != 0:
                                document.add_paragraph(
                                    "Dodatki.: " + str(el['extra_price'])+"zł")
                            if el['cake_info'] != "":
                                cake = el['cake_info']
                                cake = cake.replace("\r\n", "")
                                cake = cake.replace("  ", "")
                                document.add_paragraph(
                                    "Ciasto: "+cake+"zł"+"\n(cena wliczona w dodatki)")
                            if el['discount'] != 0:
                                document.add_paragraph(
                                    "Rabat.: " + str(el['discount'])+"%")
                            if el['change_topps_info'] != "":
                                changes = el['change_topps_info']
                                changes = changes.replace("\r\n", "")
                                changes = changes.replace("  ", "")
                                document.add_paragraph(
                                    "Zmiany lewa: "+str(changes))
                            if el['change_topps_info_other_side'] != "":
                                changes = el['change_topps_info_other_side']
                                changes = changes.replace("\r\n", "")
                                changes = changes.replace("  ", "")
                                document.add_paragraph(
                                    "Zmiany prawa: "+str(changes))
                            if el['sauces_free'] != "":
                                sauces_free = el['sauces_free']
                                document.add_paragraph(
                                    "Sosy: "+str(sauces_free))
                            if el['sauces_pay'] != "":
                                sauces_pay = el['sauces_pay']
                                sauces_pay = sauces_pay.replace("\r\n", "")
                                sauces_pay = sauces_pay.replace("  ", "")
                                document.add_paragraph(
                                    "Sosy płatne: "+str(sauces_pay))
                            if el['info'] != "":
                                document.add_paragraph(
                                    "Info.: " + str(el['info']))
                            # document.add_paragraph(
                            #     "Cena: "+str(el['total_price']))
                            document.add_paragraph("---------------- ")

                # document.save("C:\\Users\\kenny_000\\Desktop\\Python\\print.docx")
                document.save("C:\\Users\\Miktel\\Desktop\\Pythony\\print.docx")
                
                try:
                    # os.startfile("C:\\Users\\kenny_000\\Desktop\\Python\\print.docx", "print")
                    os.startfile("print.docx", "print")
                    print('Drukowanie')
                except:
                    print('Brak drukarki')
    time.sleep(1)
